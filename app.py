# app.py
from flask import Flask, request, jsonify, Response, send_file
from flask_cors import CORS
import pymupdf4llm
import pymupdf
import os
import logging
import traceback
import json
import time
from datetime import datetime
from threading import Thread
import uuid
import sys
from io import StringIO, BytesIO
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
import html2text

app = Flask(__name__)
CORS(app)

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Store conversion progress
conversion_progress = {}

class ProgressCapture:
    """Capture progress output from pymupdf4llm"""
    def __init__(self, conversion_id, total_pages):
        self.conversion_id = conversion_id
        self.total_pages = total_pages
        self.original_stdout = sys.stdout
        self.original_stderr = sys.stderr
        
    def __enter__(self):
        sys.stdout = self
        sys.stderr = self
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        sys.stdout = self.original_stdout
        sys.stderr = self.original_stderr
        
    def write(self, text):
        # Write to original stdout/stderr
        self.original_stdout.write(text)
        self.original_stdout.flush()
        
        # Parse progress from pymupdf4llm output
        if self.conversion_id in conversion_progress:
            # Look for progress patterns like "[====    ] (5/26)" or "Processing page 5 of 26"
            progress_match = re.search(r'\(\s*(\d+)/(\d+)\)', text)
            if progress_match:
                current_page = int(progress_match.group(1))
                total_pages = int(progress_match.group(2))
                
                # Calculate progress percentage (reserve 10% for finalization)
                progress_percent = int((current_page / total_pages) * 85) + 10
                
                conversion_progress[self.conversion_id].update({
                    'progress': progress_percent,
                    'stage': f'Processing page {current_page} of {total_pages}...',
                    'current_page': current_page,
                    'total_pages': total_pages
                })
                
    def flush(self):
        self.original_stdout.flush()

def convert_pdf_with_progress(temp_path, conversion_id, filename):
    """Convert PDF with real progress tracking"""
    try:
        # Get PDF metadata first
        doc = pymupdf.open(temp_path)
        total_pages = len(doc)
        file_size = os.path.getsize(temp_path)
        doc.close()
        
        # Update progress: Starting conversion
        conversion_progress[conversion_id] = {
            'progress': 0,
            'stage': 'Starting conversion...',
            'total_pages': total_pages,
            'current_page': 0,
            'filename': filename,
            'file_size': file_size,
            'status': 'processing'
        }
        
        # Start conversion with progress tracking
        logger.info(f'Starting PDF conversion for {filename}')
        
        # Update progress: Converting to markdown
        conversion_progress[conversion_id].update({
            'progress': 5,
            'stage': 'Initializing conversion...'
        })
        
        # Capture progress output from pymupdf4llm
        with ProgressCapture(conversion_id, total_pages):
            # Actual conversion - this is where the real work happens
            markdown = pymupdf4llm.to_markdown(temp_path)
        
        # Update progress: Finalizing
        conversion_progress[conversion_id].update({
            'progress': 95,
            'stage': 'Finalizing conversion...'
        })
        
        time.sleep(0.5)  # Brief pause for finalization
        
        # Format file size
        def format_file_size(size_bytes):
            if size_bytes < 1024:
                return f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f} KB"
            else:
                return f"{size_bytes / (1024 * 1024):.1f} MB"
        
        # Complete conversion
        result = {
            'markdown': markdown,
            'filename': filename,
            'fileSize': format_file_size(file_size),
            'pageCount': total_pages,
            'timestamp': datetime.now().isoformat(),
            'success': True
        }
        
        conversion_progress[conversion_id].update({
            'progress': 100,
            'stage': 'Conversion complete!',
            'status': 'completed',
            'result': result
        })
        
        logger.info('Conversion successful')
        
    except Exception as e:
        logger.error(f'Conversion error: {str(e)}')
        logger.error(traceback.format_exc())
        conversion_progress[conversion_id] = {
            'progress': 0,
            'stage': f'Error: {str(e)}',
            'status': 'error',
            'error': str(e)
        }

@app.route('/convert', methods=['POST'])
def convert():
    try:
        if 'pdf' not in request.files:
            logger.error('No file in request')
            return jsonify({'error': 'No file uploaded'}), 400
            
        file = request.files['pdf']
        if file.filename == '':
            logger.error('Empty filename')
            return jsonify({'error': 'No file selected'}), 400

        # Generate unique conversion ID
        conversion_id = str(uuid.uuid4())
        
        # Save uploaded file temporarily
        temp_path = os.path.abspath(f'temp_{conversion_id}.pdf')
        logger.info(f'Saving file to {temp_path}')
        file.save(temp_path)
        
        # Start conversion in background thread
        thread = Thread(target=convert_pdf_with_progress, args=(temp_path, conversion_id, file.filename))
        thread.start()
        
        # Return conversion ID for progress tracking
        return jsonify({
            'conversion_id': conversion_id,
            'message': 'Conversion started',
            'success': True
        })
        
    except Exception as e:
        logger.error(f'Server error: {str(e)}')
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}', 'success': False}), 500

@app.route('/progress/<conversion_id>', methods=['GET'])
def get_progress(conversion_id):
    """Get conversion progress for a specific conversion ID"""
    try:
        if conversion_id not in conversion_progress:
            return jsonify({'error': 'Conversion not found'}), 404
        
        progress_data = conversion_progress[conversion_id].copy()
        
        # Clean up completed or errored conversions after sending response
        if progress_data.get('status') in ['completed', 'error']:
            # Clean up temp file
            temp_path = os.path.abspath(f'temp_{conversion_id}.pdf')
            if os.path.exists(temp_path):
                os.remove(temp_path)
                logger.info(f'Temp file removed: {temp_path}')
            
            # Remove from progress tracking after a delay to allow final fetch
            def cleanup_progress():
                time.sleep(5)  # Wait 5 seconds before cleanup
                if conversion_id in conversion_progress:
                    del conversion_progress[conversion_id]
            
            Thread(target=cleanup_progress).start()
        
        return jsonify(progress_data)
        
    except Exception as e:
        logger.error(f'Progress error: {str(e)}')
        return jsonify({'error': f'Progress error: {str(e)}'}), 500

def markdown_to_docx(markdown_text, filename="document"):
    """Convert markdown text to a Word document"""
    try:
        # Create a new Document
        doc = Document()
        
        # Split markdown into paragraphs (separated by double newlines)
        paragraphs = re.split(r'\n\s*\n', markdown_text)
        
        for paragraph_text in paragraphs:
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
                
            # Split paragraph into lines
            lines = paragraph_text.split('\n')
            
            # Check if this is a special block (header, code, list, etc.)
            first_line = lines[0].strip()
            
            # Handle headers
            if first_line.startswith('#'):
                header_level = len(first_line) - len(first_line.lstrip('#'))
                header_text = first_line.lstrip('# ').strip()
                
                if header_level == 1:
                    doc.add_heading(header_text, level=1)
                elif header_level == 2:
                    doc.add_heading(header_text, level=2)
                elif header_level == 3:
                    doc.add_heading(header_text, level=3)
                else:
                    doc.add_heading(header_text, level=4)
                    
            # Handle code blocks
            elif first_line.startswith('```'):
                # Find all lines between ``` markers
                code_lines = []
                in_code = False
                for line in lines:
                    if line.strip().startswith('```'):
                        if in_code:
                            break
                        in_code = True
                        continue
                    if in_code:
                        code_lines.append(line)
                
                # Add code block as a paragraph with monospace font
                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_para.style = 'Normal'
                
            # Handle bullet points
            elif first_line.startswith('- ') or first_line.startswith('* '):
                for line in lines:
                    line = line.strip()
                    if line.startswith('- ') or line.startswith('* '):
                        bullet_text = line[2:].strip()
                        doc.add_paragraph(bullet_text, style='List Bullet')
                        
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', first_line):
                for line in lines:
                    line = line.strip()
                    if re.match(r'^\d+\.\s', line):
                        list_text = re.sub(r'^\d+\.\s', '', line)
                        doc.add_paragraph(list_text, style='List Number')
                        
            # Handle blockquotes
            elif first_line.startswith('>'):
                quote_lines = []
                for line in lines:
                    line = line.strip()
                    if line.startswith('>'):
                        quote_lines.append(line[1:].strip())
                    else:
                        quote_lines.append(line)
                
                quote_text = ' '.join(quote_lines)
                quote_para = doc.add_paragraph(quote_text)
                quote_para.style = 'Quote'
                
            # Handle regular paragraphs
            else:
                # Join all lines in the paragraph with spaces (single line breaks become spaces)
                full_text = ' '.join(line.strip() for line in lines if line.strip())
                
                # Process inline formatting
                para = doc.add_paragraph()
                
                # Simple bold and italic processing
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', full_text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Inline code
                        run = para.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        # Regular text
                        if part:
                            para.add_run(part)
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
        
    except Exception as e:
        logger.error(f'Error converting markdown to docx: {str(e)}')
        raise e

@app.route('/convert-markdown-to-word', methods=['POST'])
def convert_markdown_to_word():
    """Convert markdown text to Word document"""
    try:
        data = request.get_json()
        
        if not data or 'markdown' not in data:
            return jsonify({'error': 'No markdown content provided'}), 400
        
        markdown_text = data['markdown']
        filename = data.get('filename', 'document')
        
        if not markdown_text.strip():
            return jsonify({'error': 'Markdown content is empty'}), 400
        
        # Convert markdown to Word document
        doc_buffer = markdown_to_docx(markdown_text, filename)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_filename = f"{filename}_{timestamp}.docx"
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=word_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f'Error in markdown to word conversion: {str(e)}')
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Conversion error: {str(e)}'}), 500

if __name__ == '__main__':
    logger.info('Starting Flask server...')
    app.run(host='0.0.0.0', port=6201)
